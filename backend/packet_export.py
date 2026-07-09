"""Render a FaithForge packet's markdown into downloadable Word/PDF bytes.

Shares the block grammar with packet_builder.markdown_to_html (headings,
bullet/numbered/checkbox lists, pipe tables, horizontal rules, paragraphs)
so exported files match what's shown on screen.
"""
import io
import re
from typing import List, Optional, Tuple

from packet_builder import _is_table_sep, _table_cells, _normalize_row

NAVY_RGB = (0x1E, 0x3A, 0x8A)
COPPER_RGB = (0xC2, 0x65, 0x2A)


def _parse_blocks(markdown: str) -> List[Tuple]:
    lines = (markdown or "").split("\n")
    blocks: List[Tuple] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if re.match(r"^```\w*$", s):
            i += 1
            continue

        diagram_match = re.match(r"^\[\[DIAGRAM:(\w+)\]\]$", s)
        if diagram_match:
            i += 1
            blocks.append(("diagram", diagram_match.group(1)))
            continue

        if s.startswith("|") and i + 1 < len(lines) and _is_table_sep(lines[i + 1]):
            header = _table_cells(s)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_normalize_row(_table_cells(lines[i]), len(header)))
                i += 1
            blocks.append(("table", header, rows))
            continue

        i += 1

        if re.match(r"^- \[[ x]\] ", s):
            blocks.append(("checkbox", s[3] == "x", s[6:]))
            continue
        if re.match(r"^[-*] ", s):
            blocks.append(("bullet", s[2:]))
            continue
        if re.match(r"^\d+\. ", s):
            blocks.append(("numbered", re.sub(r"^\d+\. ", "", s)))
            continue

        if s.startswith("#### "):
            blocks.append(("h4", s[5:]))
        elif s.startswith("### "):
            blocks.append(("h3", s[4:]))
        elif s.startswith("## "):
            blocks.append(("h2", s[3:]))
        elif s.startswith("# "):
            blocks.append(("h1", s[2:]))
        elif s == "---":
            blocks.append(("hr",))
        elif s.startswith("*") and s.endswith("*") and len(s) > 2 and not s.startswith("**"):
            blocks.append(("italic", s.strip("*")))
        elif s:
            blocks.append(("p", s))
        else:
            blocks.append(("blank",))
    return blocks


_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")


def _split_inline(text: str) -> List[Tuple[str, bool, bool]]:
    """Split text into (segment, is_bold, is_italic) runs on **bold**/*italic* markers.

    Mirrors packet_builder._inline_md's precedence: bold (**) matches before
    single-star italic at each position.
    """
    runs = []
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], False, False))
        if m.group(1) is not None:
            runs.append((m.group(1), True, False))
        else:
            runs.append((m.group(2), False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return runs or [("", False, False)]


def _strip_bold(text: str) -> str:
    """Strip ** and * markup, leaving plain text (used for table cells)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


# ─── DOCX export ─────────────────────────────────────────────────────────────

def markdown_to_docx_bytes(markdown: str, title: str = "FaithForge Proposal", client_name: Optional[str] = None,
                            plan: Optional[dict] = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime

    doc = Document()
    doc.core_properties.title = title
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    navy = RGBColor(*NAVY_RGB)
    copper = RGBColor(*COPPER_RGB)

    # ── Cover page ────────────────────────────────────────────────────────
    company_p = doc.add_paragraph()
    company_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company_p.add_run("FAITHFORGE TECHNOLOGIES & CONSULTING")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = navy

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = navy

    if client_name:
        client_p = doc.add_paragraph()
        client_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = client_p.add_run(f"Prepared for {client_name}")
        run.font.size = Pt(13)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(datetime.now().strftime("%B %d, %Y"))
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    def add_inline_runs(paragraph, text):
        for segment, bold, italic in _split_inline(text):
            if not segment:
                continue
            run = paragraph.add_run(segment)
            run.bold = bold
            run.italic = italic

    for block in _parse_blocks(markdown):
        kind = block[0]
        if kind == "h1":
            p = doc.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = navy
        elif kind == "h2":
            p = doc.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = navy
        elif kind == "h3":
            p = doc.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = copper
        elif kind == "h4":
            p = doc.add_paragraph()
            run = p.add_run(block[1])
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
        elif kind == "hr":
            p = doc.add_paragraph("_" * 60)
            p.runs[0].font.color.rgb = copper
        elif kind == "table":
            header, rows = block[1], block[2]
            if not header:
                continue
            table = doc.add_table(rows=1, cols=len(header))
            try:
                table.style = "Light Grid Accent 1"
            except KeyError:
                pass
            for cell, text in zip(table.rows[0].cells, header):
                run = cell.paragraphs[0].add_run(_strip_bold(text))
                run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    cell.paragraphs[0].add_run(_strip_bold(text))
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, block[1])
        elif kind == "numbered":
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, block[1])
        elif kind == "checkbox":
            checked, text = block[1], block[2]
            p = doc.add_paragraph(style="List Bullet")
            box = "☑ " if checked else "☐ "
            add_inline_runs(p, box + text)
        elif kind == "italic":
            p = doc.add_paragraph()
            run = p.add_run(block[1])
            run.italic = True
        elif kind == "p":
            p = doc.add_paragraph()
            add_inline_runs(p, block[1])
        elif kind == "diagram":
            if plan:
                try:
                    from diagrams import render_diagram, png_size, fit_size_in
                    png = render_diagram(block[1], plan)
                    if png:
                        iw, ih = png_size(png)
                        display_w, _ = fit_size_in(iw, ih, 6.2, 9.0)
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(io.BytesIO(png), width=Inches(display_w))
                except Exception:
                    pass
        # blank: paragraph spacing already provides separation

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── PDF export ──────────────────────────────────────────────────────────────

_SINGLE_STAR_RE = re.compile(r"(?<!\*)\*(?!\*)([^*]+?)(?<!\*)\*(?!\*)")


def _fpdf_markdown(text: str) -> str:
    """fpdf2's native markdown uses **bold**/__italic__; our source uses
    **bold**/single *italic* — convert single-star italics for fpdf2."""
    return _SINGLE_STAR_RE.sub(r"__\1__", _pdf_safe(text))


_UNICODE_TO_ASCII = {
    "—": "-",    # em dash —
    "–": "-",    # en dash –
    "‘": "'",    # left single quote '
    "’": "'",    # right single quote '
    "“": '"',    # left double quote "
    "”": '"',    # right double quote "
    "…": "...",  # ellipsis …
    " ": " ",    # non-breaking space
    "•": "-",    # bullet •
    "☑": "[x]",  # checked box ☑
    "☐": "[ ]",  # empty box ☐
}


def _pdf_safe(text: str) -> str:
    """fpdf2's core Helvetica font is Latin-1 only. AI-generated prose
    routinely uses em-dashes/smart quotes that would otherwise crash
    rendering with FPDFUnicodeEncodingException — map the common ones to
    ASCII, then replace anything still out of range."""
    if not text:
        return text
    for uni, ascii_eq in _UNICODE_TO_ASCII.items():
        text = text.replace(uni, ascii_eq)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _column_fractions(header: List[str], rows: List[List[str]]) -> List[float]:
    """Proportion each column's width to its content length instead of splitting
    evenly — equal-width columns caused long cells to overflow past their border
    into neighboring columns (fpdf2's cell() does not wrap or clip text)."""
    n = len(header)
    if n == 0:
        return []
    avg_len = []
    for i in range(n):
        lengths = [len(_strip_bold(header[i]))]
        for row in rows[:40]:
            if i < len(row):
                lengths.append(len(_strip_bold(row[i])))
        avg_len.append((sum(lengths) / len(lengths)) if lengths else 1)
    total = sum(avg_len) or 1
    min_frac = 0.14
    fracs = [max(min_frac, l / total) for l in avg_len]
    fsum = sum(fracs) or 1
    return [f / fsum for f in fracs]


def _draw_cover_page(pdf, title: str, client_name: Optional[str] = None) -> None:
    title = _pdf_safe(title)
    client_name = _pdf_safe(client_name) if client_name else client_name
    from datetime import datetime

    pdf.add_page()
    pdf.set_fill_color(*NAVY_RGB)
    pdf.rect(0, 0, pdf.w, 65, style="F")
    pdf.set_xy(0, 24)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(pdf.w, 12, "FAITHFORGE", align="C")
    pdf.set_xy(0, 42)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(pdf.w, 8, "TECHNOLOGIES & CONSULTING", align="C")

    pdf.set_xy(pdf.l_margin, 100)
    pdf.set_text_color(*NAVY_RGB)
    pdf.set_font("Helvetica", "B", 19)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 10, title, align="C")

    if client_name:
        pdf.ln(4)
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "", 13)
        pdf.set_text_color(70, 70, 70)
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 8, f"Prepared for {client_name}", align="C")

    pdf.ln(6)
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(130, 130, 130)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 6, datetime.now().strftime("%B %d, %Y"), align="C")

    pdf.set_y(-45)
    pdf.set_draw_color(*COPPER_RGB)
    pdf.set_line_width(0.6)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(5)
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(*NAVY_RGB)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 6, "Execution Without Compromise", align="C")
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.multi_cell(
        pdf.w - pdf.l_margin - pdf.r_margin, 5,
        "410-862-2975  |  info@faithforgetech.com  |  www.faithforgetech.com", align="C",
    )


def _place_diagram_pdf(pdf, key: str, plan: dict, usable_width: float) -> None:
    """Render and embed a diagram image at the current cursor position,
    page-breaking first if it wouldn't fit on the remaining page. Sized at
    its natural resolution (capped to the page) rather than always stretched
    to full content width, so small diagrams aren't blown up."""
    from diagrams import render_diagram, png_size, fit_size_mm
    png = render_diagram(key, plan)
    if not png:
        return
    try:
        iw, ih = png_size(png)
    except Exception:
        return
    display_w, display_h = fit_size_mm(iw, ih, usable_width, 200.0)
    if pdf.get_y() + display_h > pdf.h - pdf.b_margin:
        pdf.add_page()
    y0 = pdf.get_y()
    x0 = pdf.l_margin + (usable_width - display_w) / 2
    pdf.image(io.BytesIO(png), x=x0, y=y0, w=display_w, h=display_h)
    pdf.set_xy(pdf.l_margin, y0 + display_h + 3)


def markdown_to_pdf_bytes(markdown: str, title: str = "FaithForge Proposal", client_name: Optional[str] = None,
                           plan: Optional[dict] = None) -> bytes:
    from fpdf import FPDF
    from fpdf.fonts import FontFace

    pdf = FPDF(format="Letter")
    pdf.set_title(title)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.set_margins(18, 18, 18)

    _draw_cover_page(pdf, title, client_name)
    pdf.add_page()

    black = (20, 20, 20)
    usable_width = pdf.w - pdf.l_margin - pdf.r_margin

    for block in _parse_blocks(markdown):
        kind = block[0]
        # fpdf2's multi_cell leaves the cursor at the right edge of the line
        # box, not the left margin — without this reset, a heading directly
        # followed by content (no blank-line block between them, which is
        # valid markdown and common in AI output) would compute a zero-width
        # line on the next multi_cell and crash with FPDFException.
        pdf.set_x(pdf.l_margin)
        if kind == "h1":
            pdf.set_text_color(*NAVY_RGB)
            pdf.set_font("Helvetica", "B", 15)
            pdf.multi_cell(0, 8, _pdf_safe(block[1]))
            pdf.set_draw_color(*COPPER_RGB)
            pdf.set_line_width(0.6)
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(3)
        elif kind == "h2":
            pdf.set_text_color(*NAVY_RGB)
            pdf.set_font("Helvetica", "B", 12)
            pdf.ln(2)
            pdf.multi_cell(0, 7, _pdf_safe(block[1]))
        elif kind == "h3":
            pdf.set_text_color(*COPPER_RGB)
            pdf.set_font("Helvetica", "B", 11)
            pdf.multi_cell(0, 6, _pdf_safe(block[1]))
        elif kind == "h4":
            pdf.set_text_color(*black)
            pdf.set_font("Helvetica", "BI", 10)
            pdf.multi_cell(0, 6, _pdf_safe(block[1]))
        elif kind == "hr":
            pdf.set_draw_color(*COPPER_RGB)
            y = pdf.get_y() + 2
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(6)
        elif kind == "table":
            header, rows = block[1], block[2]
            if not header:
                continue
            col_fracs = _column_fractions(header, rows)
            pdf.set_font("Helvetica", "", 9)
            pdf.set_draw_color(*COPPER_RGB)
            pdf.set_line_width(0.3)
            # fpdf2's table() snapshots whatever fill/text color is currently
            # set as its base cell style — reset here or the color from the
            # preceding heading block (e.g. navy h2 text) bleeds into every
            # body cell instead of plain black.
            pdf.set_fill_color(0, 0, 0)
            pdf.set_text_color(*black)
            head_style = FontFace(emphasis="BOLD", color=(255, 255, 255), fill_color=NAVY_RGB)
            with pdf.table(
                col_widths=col_fracs,
                headings_style=head_style,
                text_align="LEFT",
                v_align="TOP",
                line_height=5,
                padding=(1.5, 2),
                borders_layout="ALL",
                align="LEFT",
            ) as pdf_table:
                header_row = pdf_table.row()
                for h in header:
                    header_row.cell(_pdf_safe(_strip_bold(h)))
                for row in rows:
                    table_row = pdf_table.row()
                    for cell_text in row:
                        table_row.cell(_pdf_safe(_strip_bold(cell_text)))
            pdf.ln(3)
        elif kind in ("bullet", "numbered", "checkbox"):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*black)
            if kind == "checkbox":
                checked, text = block[1], block[2]
                prefix = "[x] " if checked else "[ ] "
            else:
                text = block[1]
                prefix = "-  "
            pdf.set_x(pdf.l_margin + 4)
            pdf.multi_cell(usable_width - 4, 5.5, _fpdf_markdown(prefix + text), markdown=True)
        elif kind == "italic":
            pdf.set_font("Helvetica", "I", 10)
            pdf.set_text_color(*black)
            pdf.multi_cell(0, 5.5, _pdf_safe(block[1]))
        elif kind == "p":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*black)
            pdf.multi_cell(0, 5.5, _fpdf_markdown(block[1]), markdown=True)
        elif kind == "diagram":
            if plan:
                _place_diagram_pdf(pdf, block[1], plan, usable_width)
        elif kind == "blank":
            pdf.ln(2)

    return bytes(pdf.output())
